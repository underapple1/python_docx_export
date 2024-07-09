# ! /usr/bin/env python3
#  -*- coding: utf-8 -*-
#
# yumizi @ 2023-06
#
# ä¸€å †å±å±±ä»£ç 
#
import json
import sys
import time
from threading import Thread
from tkinter.filedialog import asksaveasfile

import docx

import log
import os
import tkinter as tk
from tkinter import filedialog
from tkinter.messagebox import askyesno, showinfo, showwarning

from gui import ApplicationGUI, ExportGUI

# ç‰ˆæœ¬
from oletools import oleobj

v = 'v 0.1.0.3'

# æ—¥å¿—å¯¹è±¡
logger = log.Logger().log_create()

# APPæ ¹ç›®å½•
app_path = os.path.dirname(os.path.abspath(__file__))
# èµ„æºåµŒå…¥
mp = ''.join(['_', 'M', 'E', 'I', 'P', 'A', 'S', 'S'])


def resource_path(file_name):
    if hasattr(sys, mp):
        return os.path.join(getattr(sys, mp), file_name)
    return os.path.join(file_name)


def file_check(file_path: str):
    if os.path.basename(file_path).startswith('~$'):
        logger.debug('æœªæ·»åŠ ï¼šç¼“å­˜æ–‡ä»¶ %s', file_path)
        return False
    if not os.path.basename(file_path).endswith('.docx'):
        logger.debug('æœªæ·»åŠ ï¼šédocxæ–‡ä»¶ %s', file_path)
        return False
    if not os.path.isfile(file_path):
        logger.debug('æœªæ·»åŠ ï¼šä¸æ˜¯æ–‡ä»¶ %s', file_path)
        return False
    return True


class Application(ApplicationGUI):
    def __init__(self, master=None, version=''):
        super().__init__(master)
        # å¯¼å‡ºåˆ—è¡¨
        self.export_list = {}
        # å¤±è´¥åˆ—è¡¨
        self.fail_list = []

        self.master = master
        self.version = version

        # å¾…å¤„ç†æ–‡ä»¶åˆ—è¡¨
        self.file_list = []
        # æ–‡ä»¶æ·»åŠ ç´¢å¼•
        self.file_index = 0

        self.main_listbox = None
        self.create_list(self.frame_list)
        self.set_combobox_text()
        self.set_checkbutton_text()
        self.register_command()

        self.entry_tips_val.set('ğŸ«…ğŸ¤´ğŸ‘¸')

        self.export_ui = None
        self.e_progress = 0
        self.e_total_task = 0
        self.e_rename = 0
        self.e_cancel = 0
        self.e_pause = 0

        self.lb_brand.configure(text='èµ„äº§å¯¼å‡º\n' + self.version + '\n@MSS')

    def set_checkbutton_text(self):
        self.che_text.set(1)
        self.che_image.set(1)
        self.che_table.set(1)
        self.che_combine.set(1)
        self.che_attachment.set(1)
        self.che_info.set(1)

    def set_combobox_text(self):
        self.cb_save_position['value'] = ('åŸæ–‡ä»¶ç›®å½•ä¸­ä»¥åŸæ–‡ä»¶å‘½åçš„å­æ–‡ä»¶å¤¹ä¸­', 'åŸæ–‡ä»¶æ‰€åœ¨æ–‡ä»¶å¤¹', 'è‡ªå®šä¹‰æ–‡ä»¶å¤¹')
        self.combobox_save_path.set(self.cb_save_position['value'][0])
        self.cb_save_position.current(0)

        self.cb_export_cover['value'] = ('åŒè·¯å¾„æ–‡ä»¶ï¼šè¦†ç›–', 'åŒè·¯å¾„æ–‡ä»¶ï¼šè‡ªåŠ¨é‡å‘½å', 'åŒè·¯å¾„æ–‡ä»¶ï¼šè·³è¿‡')
        self.combobox_export_cover.set(self.cb_export_cover['value'][0])

        self.cb_name_part1['value'] = ('|è‡ªå¢ç¼–å·|', '|è¿æ¥ç¬¦|', '|åŸæ–‡ä»¶å|', '|åç¼€å|')
        self.cb_name_part2['value'] = ('|è¿æ¥ç¬¦|', '|åŸæ–‡ä»¶å|', '|åç¼€å|', '|è‡ªå¢ç¼–å·|')
        self.cb_name_part3['value'] = ('|åŸæ–‡ä»¶å|', '|åç¼€å|', '|è‡ªå¢ç¼–å·|', '|è¿æ¥ç¬¦|')
        self.cb_name_part4['value'] = ('|åç¼€å|', '|è‡ªå¢ç¼–å·|', '|è¿æ¥ç¬¦|', '|åŸæ–‡ä»¶å|')

        self.combobox_name1.set(self.cb_name_part1['value'][0])
        self.combobox_name2.set(self.cb_name_part2['value'][0])
        self.combobox_name3.set(self.cb_name_part3['value'][0])
        self.combobox_name4.set(self.cb_name_part4['value'][0])
        self.cb_name_part1.current(0)
        self.cb_name_part2.current(0)
        self.cb_name_part3.current(0)
        self.cb_name_part4.current(0)

    def create_list(self, frame):
        # ä¸€ä¸ªåˆ—è¡¨
        list_frame = tk.Frame(frame)
        scroll_h_bar = tk.Scrollbar(list_frame, orient=tk.HORIZONTAL)  # æ°´å¹³æ»šåŠ¨æ¡ç»„ä»¶
        scroll_v_bar = tk.Scrollbar(list_frame, orient=tk.VERTICAL, )  # å‚ç›´æ»šåŠ¨æ¡ç»„ä»¶
        self.main_listbox = tk.Listbox(list_frame,
                                       width=60, height=20,
                                       selectmode=tk.MULTIPLE,
                                       yscrollcommand=scroll_v_bar.set,
                                       xscrollcommand=scroll_h_bar.set)

        scroll_v_bar.pack(side=tk.RIGHT, fill=tk.Y)  # è®¾ç½®å‚ç›´æ»šåŠ¨æ¡æ˜¾ç¤ºçš„ä½ç½®
        scroll_v_bar.config(command=self.main_listbox.yview)  # è®¾ç½®Scrollbarç»„ä»¶çš„commandé€‰é¡¹ä¸ºè¯¥ç»„ä»¶çš„yview()æ–¹æ³•
        scroll_h_bar.pack(side=tk.BOTTOM, fill=tk.X)  # è®¾ç½®æ°´å¹³æ»šåŠ¨æ¡æ˜¾ç¤ºçš„ä½ç½®
        scroll_h_bar.config(command=self.main_listbox.xview)  # è®¾ç½®Scrollbarç»„ä»¶çš„commandé€‰é¡¹ä¸ºè¯¥ç»„ä»¶çš„xview()æ–¹æ³•
        self.main_listbox.place(relx=0, rely=0, relheight=0.95, relwidth=0.952, bordermode='ignore')
        list_frame.place(relx=0, rely=0, relheight=1, relwidth=1, bordermode='ignore')

    def register_command(self):
        # é€‰æ‹©æ–‡ä»¶
        self.btn_import_files.bind('<Button>', self.choose_file)
        # é€‰æ‹©æ–‡ä»¶å¤¹
        self.btn_import_dir.bind('<Button>', self.choose_dir)
        # åˆ é™¤åˆ—è¡¨é€‰ä¸­é¡¹
        self.btn_delete_list_items.bind('<Button>', self.remove_list_item)
        # åˆ é™¤æ‰€æœ‰é¡¹
        self.btn_delete_list_all.bind('<Button>', self.remove_list_all)
        # é€‰æ‹©æ–‡ä»¶å¤¹
        self.btn_choose_position.bind('<Button>', self.choose_export_dir)
        # å¯¼å‡ºæŒ‰é’®
        self.btn_export.bind('<Button-1>', self.export)
        # ---------------------------------------------------------------------------
        self.bind_cb_evt(self.cb_name_part1, self.name_tips)
        self.bind_cb_evt(self.cb_name_part2, self.name_tips)
        self.bind_cb_evt(self.cb_name_part3, self.name_tips)
        self.bind_cb_evt(self.cb_name_part4, self.name_tips)
        self.bind_cb_evt(self.cb_export_cover, self.export_tips_cover)
        # ---------------------------------------------------------------------------
        self.bind_cb_evt(self.cb_save_position, self.dir_tips)
        self.bind_cb_evt(self.entry_save_position, self.dir_tips)
        # ---------------------------------------------------------------------------
        self.cb_delete_raw_file.bind('<Button>', self.delete_tips)
        # ---------------------------------------------------------------------------
        self.cb_export_type_text.bind('<Button>', self.export_tips_text)
        self.cb_export_type_image.bind('<Button>', self.export_tips_image)
        self.cb_export_type_attachment.bind('<Button>', self.export_tips_attachment)
        self.cb_export_type_table.bind('<Button>', self.export_tips_table)
        self.cb_export_type_combine_text_table.bind('<Button>', self.export_tips_combine_text_table)
        self.cb_export_type_info.bind('<Button>', self.export_tips_info)

    @staticmethod
    def bind_cb_evt(cb, evt):
        cb.bind('<Button>', evt)
        cb.bind('<<ComboboxSelected>>', evt)
        cb.bind('<space>', evt)
        cb.bind('<Return>', evt)
        cb.bind('<Key>', evt)

    def export_tips_cover(self, evt):
        if not evt:
            return
        way = self.combobox_export_cover.get()
        if way == self.cb_export_cover['value'][0]:
            # è¦†ç›–
            self.entry_tips_val.set('å¯¼å‡ºæ—¶ç›¸åŒè·¯å¾„çš„æ–‡ä»¶ä¼šè¢«è¦†ç›–æ‰')
        elif way == self.cb_export_cover['value'][1]:
            # é‡å‘½å
            self.entry_tips_val.set('å¯¼å‡ºæ—¶è‹¥æ–‡ä»¶å·²å­˜åœ¨åˆ™è‡ªåŠ¨é‡å‘½å')
        elif way == self.cb_export_cover['value'][2]:
            # è·³è¿‡
            self.entry_tips_val.set('å¯¼å‡ºæ—¶è‹¥æ–‡ä»¶å·²å­˜åœ¨åˆ™è·³è¿‡')
        else:
            # æ„å¤–
            self.entry_tips_val.set('è¿™æ˜¯ä»€ä¹ˆæƒ…å†µï¼Ÿï¼Ÿï¼')

    def export_tips_image(self, evt):
        if not evt:
            return
        # è¿™ä¸ªå€¼æ˜¯ç‚¹å‡»ä¹‹å‰çš„å€¼ï¼Œ
        if not self.che_image.get():
            self.entry_tips_val.set('å¯¼å‡ºWordä¸­çš„å›¾ç‰‡')
        else:
            self.entry_tips_val.set('ä¸å¯¼å‡ºWordä¸­çš„å›¾ç‰‡')

    def export_tips_attachment(self, evt):
        if not evt:
            return
        # è¿™ä¸ªå€¼æ˜¯ç‚¹å‡»ä¹‹å‰çš„å€¼ï¼Œ
        if not self.che_attachment.get():
            self.entry_tips_val.set('å¯¼å‡ºWordä¸­çš„é™„ä»¶')
        else:
            self.entry_tips_val.set('ä¸å¯¼å‡ºWordä¸­çš„é™„ä»¶')

    def export_tips_table(self, evt):
        if not evt:
            return
        # è¿™ä¸ªå€¼æ˜¯ç‚¹å‡»ä¹‹å‰çš„å€¼ï¼Œ
        if not self.che_table.get():
            self.entry_tips_val.set('å¯¼å‡ºWordè¡¨æ ¼ä¸­çš„æ–‡å­—')
        else:
            self.entry_tips_val.set('ä¸å¯¼å‡ºWordè¡¨æ ¼ä¸­çš„æ–‡å­—')

    def export_tips_info(self, evt):
        if not evt:
            return
        # è¿™ä¸ªå€¼æ˜¯ç‚¹å‡»ä¹‹å‰çš„å€¼ï¼Œ
        if not self.che_info.get():
            self.entry_tips_val.set('å¯¼å‡ºWordæ–‡æ¡£çš„ä¿¡æ¯ï¼ˆä½œè€…ã€ä¿®æ”¹æ—¶é—´å•¥çš„ï¼‰')
        else:
            self.entry_tips_val.set('ä¸å¯¼å‡ºWordæ–‡æ¡£çš„ä¿¡æ¯')

    def export_tips_combine_text_table(self, evt):
        if not evt:
            return
        # è¿™ä¸ªå€¼æ˜¯ç‚¹å‡»ä¹‹å‰çš„å€¼ï¼Œ
        if not self.che_combine.get():
            self.entry_tips_val.set('åˆå¹¶å¯¼å‡ºWordä¸­çš„æ™®é€šæ–‡å­—å’Œè¡¨æ ¼æ–‡å­—')
        else:
            self.entry_tips_val.set('åˆ†åˆ«å¯¼å‡ºWordä¸­çš„æ™®é€šæ–‡å­—å’Œè¡¨æ ¼æ–‡å­—')

    def export_tips_text(self, evt):
        if not evt:
            return
        # è¿™ä¸ªå€¼æ˜¯ç‚¹å‡»ä¹‹å‰çš„å€¼ï¼Œ
        if not self.che_text.get():
            self.entry_tips_val.set('å¯¼å‡ºWordä¸­çš„æ™®é€šæ–‡å­—')
        else:
            self.entry_tips_val.set('ä¸å¯¼å‡ºWordä¸­çš„æ™®é€šæ–‡å­—')

    def delete_tips(self, evt):
        if not evt:
            return
        # è¿™ä¸ªå€¼æ˜¯ç‚¹å‡»ä¹‹å‰çš„å€¼ï¼Œ
        if not self.che_delete_raw.get():
            self.entry_tips_val.set('å¯¼å‡ºæˆåŠŸåç«‹å³åˆ é™¤åŸæ–‡ä»¶')
        else:
            self.entry_tips_val.set('å¯¼å‡ºæˆåŠŸåä¿ç•™åŸæ–‡ä»¶')

    def dir_tips(self, evt):
        if not evt:
            return
        way = self.combobox_save_path.get()
        if way == self.cb_save_position['value'][0]:
            # åŸæ–‡ä»¶å­ç›®å½•
            self.entry_tips_val.set('ä¼šåœ¨åŸæ–‡ä»¶æ‰€åœ¨ç›®å½•åˆ›å»ºä¸€ä¸ªåŒåå­æ–‡ä»¶å¤¹ä»¥å­˜å‚¨')
        elif way == self.cb_save_position['value'][1]:
            # åŸæ–‡ä»¶ç›®å½•
            self.entry_tips_val.set('ä¿å­˜ä½ç½®ä¸åŸæ–‡ä»¶åœ¨åŒä¸€æ–‡ä»¶å¤¹')
        else:
            export_dir = self.entry_save_position_val.get()
            # æŒ‡å®šç›®å½•
            self.entry_tips_val.set('ä¿å­˜ä½ç½®ï¼š' + export_dir)

    def name_tips(self, evt):
        if not evt:
            return
        name = self.combobox_name1.get()
        name += self.combobox_name2.get()
        name += self.combobox_name3.get()
        name += self.combobox_name4.get()

        name = name.replace('|è‡ªå¢ç¼–å·|', 'ç¼–å·') \
            .replace('|è¿æ¥ç¬¦|', ' - ') \
            .replace('|åŸæ–‡ä»¶å|', 'åŸæ–‡ä»¶å') \
            .replace('|åç¼€å|', '.åç¼€') \
            .strip()

        self.entry_tips_val.set('æ–‡ä»¶åæ ¼å¼ï¼š%s' % name)

    def choose_export_dir(self, evt):
        if not evt:
            return
        directory = filedialog.askdirectory()
        if directory:
            self.entry_save_position_val.set(directory)
        self.dir_tips(evt)  # é€‰æ‹©æ–‡ä»¶å¤¹åæ˜¾ç¤ºä¸€ä¸‹

    def remove_list_all(self, evt):
        if not evt:
            return
        if not askyesno('åˆ é™¤æ‰€æœ‰é¡¹', 'æ˜¯å¦æ¸…ç©ºå·²å¯¼å…¥åˆ—è¡¨ï¼Ÿ'):
            return

        if isinstance(self.main_listbox, tk.Listbox):
            self.main_listbox.delete(0, 'end')

        # å¾…å¤„ç†æ–‡ä»¶åˆ—è¡¨
        self.file_list = []
        # æ–‡ä»¶æ·»åŠ ç´¢å¼•
        self.file_index = 0

        self.entry_tips_val.set('å·²æ¸…ç©ºåˆ—è¡¨')

    def remove_list_item(self, evt):
        if not evt:
            return
        if len(self.main_listbox.curselection()) == 0:
            showwarning('é”™è¯¯', 'æ²¡æœ‰é€‰ä¸­ä»»ä½•æ¡ç›®')
            self.entry_tips_val.set('æ²¡æœ‰ä»åˆ—è¡¨ä¸­ç§»é™¤ä»»ä½•æ–‡ä»¶')
        else:
            success = 0
            remove_index_list = []
            for index in self.main_listbox.curselection():
                remove_index_list.append(index)
                self.main_listbox.select_clear(index)

            # éœ€è¦ä»åé¢å¼€å§‹åˆ é™¤
            remove_index_list.sort(reverse=True)

            remove_file_list = []
            for index in remove_index_list:
                item_text = self.main_listbox.get(index)
                if item_text:
                    file_path = str(item_text).split('|-', 1)[1]
                    remove_file_list.append(file_path)
                    self.main_listbox.delete(index)

            for file_path in remove_file_list:
                self.file_list.remove(file_path)
                success += 1

            self.entry_tips_val.set('å·²ä»åˆ—è¡¨ä¸­ç§»é™¤{0}ä¸ªæ–‡ä»¶'.format(success))

    def add_file_list(self, files, is_choose_son=False):
        file_list = []

        # éå†ç›®å½•ï¼Œæ‹¿åˆ°æ–‡ä»¶åˆ—è¡¨
        if isinstance(files, str) and os.path.isdir(files):
            logger.debug('æ·»åŠ æ–‡ä»¶å¤¹')
            directory = files
            logger.debug('æ·»åŠ å­æ–‡ä»¶å¤¹ä¸­çš„å†…å®¹ï¼š%s é€‰æ‹©ç›®å½•ï¼š%s', is_choose_son, directory)
            for root, dirs, files in os.walk(directory):  # éå†ç›®å½•
                if directory != root and not is_choose_son:  # è·³è¿‡å­æ–‡ä»¶å¤¹
                    continue
                logger.debug('æ·»åŠ ç›®å½•ï¼š%s', root)
                for file in files:  # éå†æ–‡ä»¶
                    file_path = os.path.join(root, file)  # æ‹¼æ¥è·¯å¾„
                    file_list.append(file_path)  # æ·»åŠ åˆ°åˆ—è¡¨ä¸­
        elif isinstance(files, tuple):
            logger.debug('æ·»åŠ æ–‡ä»¶')
            file_list = files
        else:
            showwarning('æç¤º', 'å¯¼å…¥æ–‡ä»¶å‚æ•°ä¸æ­£ç¡®')

        success = 0

        # åºå·å‰é¢è¡¥ 0
        bit = len(str(len(file_list)))
        if bit < 4:
            bit = 4

        # æ·»åŠ åˆ°åˆ—è¡¨è§†å›¾ä¸­
        for file in file_list:
            if file_check(file):
                # å·²æ·»åŠ ï¼Œä¸éœ€è¦å†æ·»åŠ 
                if file in self.file_list:
                    continue
                self.file_index += 1
                logger.debug('æ·»åŠ æ–‡ä»¶ï¼š%s', file)
                self.main_listbox.insert(0, str(self.file_index).rjust(bit, '0') + '|-' + file)  # ä»æœ€åä¸€ä¸ªä½ç½®å¼€å§‹åŠ å…¥å€¼
                self.file_list.append(file)
                success += 1

        tip_str = 'æˆåŠŸæ·»åŠ {0}ä¸ªdocxæ–‡ä»¶ï¼Œå¤±è´¥{1}ä¸ª'.format(success, len(file_list) - success)
        self.entry_tips_val.set(tip_str)
        logger.debug(tip_str)
        showinfo('æ·»åŠ ç»“æœ', 'æˆåŠŸæ·»åŠ {0}ä¸ªdocxæ–‡ä»¶'.format(success))

    def choose_file(self, evt):
        if not evt:
            return
        self.entry_tips_val.set('æ­£åœ¨å¯¼å…¥ï¼Œè¯·ç­‰å¾….........')
        files_tuple = filedialog \
            .askopenfilename(title='è¯·é€‰æ‹©docxæ–‡ä»¶', filetypes=[('Word', '.docx')],
                             defaultextension='.docx',
                             multiple=True)
        if files_tuple:
            Thread(target=self.add_file_list, args=(files_tuple,)).start()

    def choose_dir(self, evt):
        """
        é€‰æ‹©æ–‡ä»¶å¤¹
        :param evt: äº‹ä»¶
        :return: None
        """
        if not evt:
            return
        self.entry_tips_val.set('æ­£åœ¨å¯¼å…¥ï¼Œè¯·ç­‰å¾….........')
        is_choose_son = askyesno('é€‰æ‹©æ–‡ä»¶å¤¹', 'é€‰æ‹©æ–‡ä»¶å¤¹æ—¶æ˜¯å¦é€‰æ‹©æ‰€æœ‰å­æ–‡ä»¶å¤¹å†…çš„æ–‡ä»¶ï¼Ÿ')
        directory = filedialog.askdirectory()
        if directory:
            Thread(target=self.add_file_list, args=(directory, is_choose_son)).start()

    def export(self, evt):
        """
        å¯¼å‡ºæŒ‰é’®äº‹ä»¶ï¼Œè´Ÿè´£æ ¡éªŒå‚æ•°ï¼Œæ•´åˆå‚æ•°
        :param evt: äº‹ä»¶
        :return: None
        """
        if not evt:
            return
        export_dir_choose = self.combobox_save_path.get()
        logger.debug('å¯¼å‡ºä½ç½®ï¼š%s', export_dir_choose)

        export_dir = self.entry_save_position_val.get()
        logger.debug('å¯¼å‡ºç›®å½•ï¼š%s', export_dir)

        if export_dir_choose == self.cb_save_position['value'][2]:
            if not (export_dir and os.path.isdir(export_dir)):
                showwarning('å¯¼å‡ºæ—¶é‡åˆ°é—®é¢˜', 'ä¿å­˜ä½ç½®æœªè®¾ç½®æ–‡ä»¶å¤¹ï¼Œè¯·è®¾ç½®ä¿å­˜æ–‡ä»¶å¤¹æˆ–è®¾ç½®å…¶ä»–ä¿å­˜ä½ç½®')
                return

        try:
            export_dir_choose = self.cb_save_position['value'].index(export_dir_choose) + 1
            logger.debug('å¯¼å‡ºæ–¹å¼ï¼š%s', export_dir_choose)
        except ValueError:
            showwarning('å‡ºé”™', 'ä¿å­˜æ–¹å¼å‚æ•°é”™è¯¯')
            return

        name = self.combobox_name1.get()
        name += self.combobox_name2.get()
        name += self.combobox_name3.get()
        name += self.combobox_name4.get()
        logger.debug('åå­—è§„åˆ™ï¼š%s', name)
        if not name.endswith('|åç¼€å|'):
            if not askyesno('æç¤º', 'æ–‡ä»¶åæœ€å¥½ä»¥â€œ|åç¼€å|â€ç»“å°¾ï¼Œå¦åˆ™å¯èƒ½å¯¼è‡´æ— æ³•è¯†åˆ«ï¼Œæ˜¯å¦ç»§ç»­ï¼Ÿ'):
                return

        export_type = []
        if self.che_text.get():
            export_type.append('text')
        if self.che_table.get():
            export_type.append('table')
        if self.che_image.get():
            export_type.append('image')
        if self.che_attachment.get():
            export_type.append('attachment')
        if self.che_combine.get():
            export_type.append('combine')
        if self.che_info.get():
            export_type.append('info')

        logger.debug('å¯¼å‡ºç±»å‹ï¼š%s', str(export_type))

        is_delete = self.che_delete_raw.get()
        logger.debug('å¯¼å‡ºååˆ é™¤åŸæ–‡ä»¶ï¼š%s', str(is_delete))

        cover = self.combobox_export_cover.get()
        if cover == self.cb_export_cover['value'][0]:
            # è¦†ç›–
            cover = 'cover'
            logger.debug('å¯¼å‡ºæ—¶ç›¸åŒè·¯å¾„çš„æ–‡ä»¶ä¼šè¢«è¦†ç›–æ‰')
        elif cover == self.cb_export_cover['value'][1]:
            # é‡å‘½å
            cover = 'rename'
            logger.debug('å¯¼å‡ºæ—¶è‹¥æ–‡ä»¶å·²å­˜åœ¨åˆ™è‡ªåŠ¨é‡å‘½å')
        elif cover == self.cb_export_cover['value'][2]:
            # è·³è¿‡
            cover = 'skip'
            logger.debug('å¯¼å‡ºæ—¶è‹¥æ–‡ä»¶å·²å­˜åœ¨åˆ™è·³è¿‡')
        else:
            # æ„å¤–
            showwarning('æç¤º', 'å‚æ•°å‡ºç°äº†æ„å¤–ï¼ŒæœªçŸ¥çš„æ–‡ä»¶è¦†ç›–ç­–ç•¥')

        parameter = {
            'way': export_dir_choose,  # ä¿å­˜æ–¹å¼
            'dir': export_dir,  # ä¿å­˜ç›®å½•
            'name': name,  # æ–‡ä»¶åæ ¼å¼
            'type': export_type,  # å¯¼å‡ºç±»å‹
            'del': bool(is_delete),  # å¯¼å‡ºååˆ é™¤åŸæ–‡ä»¶
            'list': self.file_list,  # å¯¼å‡ºæ–‡ä»¶
            'cover': cover,  # è¦†ç›–ç­–ç•¥
        }

        # å¼€å§‹å¯¼å‡º
        self.start_export(parameter)

    # ====================================================================================

    def start_export(self, parameter):
        """
        åˆå§‹åŒ–å¯¼å‡ºç•Œé¢ï¼Œå¯åŠ¨å¯¼å‡ºçº¿ç¨‹
        :param parameter: å¯¼å‡ºå‚æ•°
        :return: None
        """
        logger.debug('å¯åŠ¨ç•Œé¢')
        # åˆå§‹åŒ–å¯¼å‡ºç•Œé¢å‚æ•°
        self.e_cancel = 0
        self.e_pause = 0
        self.e_progress = 0
        self.e_total_task = len(parameter.get('list'))
        # åˆå§‹åŒ–ç•Œé¢
        self.export_ui = ExportGUI(self.master)
        # ä¸­é—´æ–‡æœ¬æ¡†åˆå§‹åŒ–
        self.export_ui.txt_d_result_show.delete('1.0', 'end')
        self.export_ui.txt_d_result_show.insert('end', 'å‡†å¤‡å¯¼å‡º......\n')
        # åˆ‡æ¢åˆ°è¿›åº¦æ¡ç•Œé¢
        self.export_ui.switch_func(None, 'æš‚åœ', 'å–æ¶ˆ', None, None)
        # é‡ç½®æç¤ºä¿¡æ¯
        self.export_ui.lb_d_tips_var.set('''å¤„ç†è¿›åº¦ï¼šâˆ%''')
        # é‡ç½®è¿›åº¦å’Œé€‰é¡¹
        self.export_ui.pd_d_main_var.set(0)  # è¿›åº¦
        self.export_ui.cb_d_all_var.set(0)  # ä¸€å¾‹å¦‚æ­¤é€‰é¡¹
        # ç»‘å®šç‚¹å‡»äº‹ä»¶
        self.export_ui.btn_d_e.bind('<Button-1>', self.e_btn_ex)
        self.export_ui.btn_d_left.bind('<Button-1>', self.e_btn_left)
        self.export_ui.btn_d_right.bind('<Button-1>', self.e_btn_right)

        Thread(target=self.run, args=(parameter,)).start()

    def show_fail(self):
        for i in range(7):
            time.sleep(0.05)
            self.export_ui.txt_d_result_show.insert(1.0, '-> \n')
        for i, file in enumerate(self.fail_list):
            time.sleep(0.01)
            self.export_ui.txt_d_result_show.insert(1.0, 'å¤±è´¥{0} -> {1}\n'.format(str(i + 1).rjust(3, '0'), file))
        show_msg = """
-->> å¤±è´¥çš„åŸå› å¯èƒ½æ˜¯æ–‡ä»¶æŸåæˆ–æ— å†…å®¹ï¼Œè¯·å°è¯•ç”¨WPSæ‰“å¼€å¹¶å¦å­˜ä¸ºdocx <<--
-->> >>>>>>  å¤åˆ¶åˆ°æ–‡ä»¶ç®¡ç†å™¨åœ°å€æ ä¸­Enterå¯ç›´æ¥æ‰“å¼€  <<<<<<< <<--
\n"""
        show_msg = list(show_msg)
        show_msg.reverse()
        for i, c in enumerate(show_msg):
            if str(c) == '\n':
                time.sleep(0.4)
            else:
                time.sleep(0.015)
            self.export_ui.txt_d_result_show.insert(1.0, str(c))

    def e_btn_ex(self, evt):
        """
        ç‚¹å‡»â€œæ‰©å±•â€æŒ‰é’®
        :param evt: äº‹ä»¶
        :return: None
        """
        if not evt:
            return
        if isinstance(self.export_ui, ExportGUI):
            btn_text = self.export_ui.btn_d_e_var.get()
            if btn_text == 'å¤±è´¥åˆ—è¡¨':
                Thread(target=self.show_fail).start()

    def e_btn_left(self, evt):
        """
        ç‚¹å‡»â€œæš‚åœâ€æŒ‰é’®
        :param evt: äº‹ä»¶
        :return: None
        """
        if not evt:
            return
        if isinstance(self.export_ui, ExportGUI):
            btn_text = self.export_ui.btn_d_left_var.get()
            if btn_text == 'æš‚åœ':
                self.e_pause = 1
                self.export_ui.btn_d_left_var.set('ç»§ç»­')
                self.show_progress(0, '.............................', 'å·²æš‚åœ')
            elif btn_text == 'ç»§ç»­':
                self.e_pause = 0
                self.export_ui.btn_d_left_var.set('æš‚åœ')
            elif btn_text == 'å¯¼å‡ºæŠ¥å‘Š':
                self.save_json(self.export_list)

    def e_btn_right(self, evt):
        """
        ç‚¹å‡»â€œå–æ¶ˆâ€æŒ‰é’®
        :param evt: äº‹ä»¶
        :return: None
        """
        if not evt:
            return
        if isinstance(self.export_ui, ExportGUI):
            btn_text = self.export_ui.btn_d_right_var.get()
            if btn_text == 'å–æ¶ˆ':
                self.e_cancel = 1
            elif btn_text == 'å…³é—­':
                self.export_ui.tf_d_title.place(relx=0.0, rely=0.0, relheight=0.0, relwidth=0.0)
                self.export_ui.tf_d_title.destroy()

    def show_progress(self, current, file, state):
        """
        æ˜¾ç¤ºè¿›åº¦
        :param current: å½“å‰åºå·
        :param file: æ–‡ä»¶
        :param state: å¯¼å‡ºçŠ¶æ€
        :return: None
        """

        # è®¡ç®—ç™¾åˆ†æ¯”
        if self.e_total_task != 0:
            current = (current + 1) / self.e_total_task * 100
        else:
            current = 0

        # ç™¾åˆ†æ¯”ä¸èƒ½æ‰ï¼ˆå•çº¿ç¨‹å¥½åƒæ²¡å¿…è¦ï¼‰
        if self.e_progress > current:
            current = self.e_progress
        else:
            self.e_progress = current

        # è®¾ç½®ç•Œé¢æ˜¾ç¤ºç™¾åˆ†æ¯”
        self.export_ui.pd_d_main_var.set(current)
        self.export_ui.lb_d_tips_var.set('å¤„ç†è¿›åº¦ï¼š{0}%'.format(current))
        # åœ¨æ–‡æœ¬æ¡†ä¸­æ˜¾ç¤ºç»†èŠ‚
        self.export_ui.txt_d_result_show.insert(1.0, '{0} -> {1}\n'.format(state, file))

    def run(self, parameter):
        print('å¼€å§‹å¤„ç†ä»»åŠ¡')

        success = 0
        sub_file_count = 0

        # å¤„ç†å‚æ•°
        file_list = parameter.get('list')
        if not isinstance(file_list, list):
            print('å‚æ•°é”™è¯¯')
            return

        # æ¸…ç©ºä¸Šæ¬¡æŠ¥å‘Š
        self.export_list = {}

        self.fail_list = []

        for index, file in enumerate(file_list):
            self.export_list[file] = []
            time.sleep(0.01)
            if self.e_cancel:
                if askyesno('å–æ¶ˆå¯¼å‡º', 'æ˜¯å¦å–æ¶ˆå¯¼å‡ºï¼Ÿï¼ˆå·²å¯¼å‡ºçš„ä¸ä¼šè¢«æ¸…ç†ï¼‰'):
                    break
            # æš‚åœ
            while self.e_pause:
                continue
            logger.debug('å¤„ç†æ–‡ä»¶ï¼š%s', file)
            try:
                sub_file_list = self.dispose(index, file, parameter)
                if isinstance(sub_file_list, bool) and not sub_file_list:
                    raise IOError('æ–‡ä»¶æ‰“å¼€å‡ºé”™')
                if not sub_file_list:
                    sub_file_list = []
            except Exception as e:
                logger.exception('å¯¼å‡ºé”™è¯¯ï¼šå¯¼å‡ºæ–‡ä»¶æ—¶å‡ºé”™ %s' % str(e))
                self.fail_list.append(file)
                self.show_progress(index, file, 'å¤±è´¥\t'.rjust(6, '-'))
                continue
            logger.debug('æˆåŠŸ %s', file)
            # ç»Ÿè®¡è®¡æ•°
            sub_count = len(sub_file_list)
            success += 1
            sub_file_count += sub_count
            # ä¿å­˜åˆ°æ€»çš„å¯¼å‡ºåˆ—è¡¨ä¸­
            self.export_list[file] = sub_file_list
            self.show_progress(index, file, 'æˆåŠŸ %s\t' % str(sub_count).rjust(2, '0'))

        if len(file_list) == 0:
            logger.debug('æ²¡æœ‰éœ€è¦å¤„ç†çš„æ–‡ä»¶')
            self.entry_tips_val.set('æ²¡æœ‰éœ€è¦å¤„ç†çš„æ–‡ä»¶')
            self.export_ui.txt_d_result_show.insert(1.0, 'æ²¡æœ‰éœ€è¦å¤„ç†çš„æ–‡ä»¶\n')
            self.export_ui.switch_func(None, None, 'å…³é—­', None, None)
            showinfo('æç¤º', 'æ²¡æœ‰éœ€è¦å¤„ç†çš„æ–‡ä»¶')
        else:
            # æ€»ç»“æç¤º
            showinfo('å¯¼å‡ºç»“æœ',
                     'å¯¼å‡ºå®Œæˆ\næˆåŠŸå¯¼å‡º%dä¸ªdocxæ–‡æ¡£\nç”Ÿæˆå­æ–‡æ¡£%dä¸ª\nå¤±è´¥%dä¸ª' % (success, sub_file_count, len(file_list) - success))

            self.entry_tips_val.set('å¯¼å‡ºå®Œæˆ')
            logger.debug('ä»»åŠ¡å®Œæˆ')
            if self.fail_list:
                self.export_ui.switch_func('å¤±è´¥åˆ—è¡¨', 'å¯¼å‡ºæŠ¥å‘Š', 'å…³é—­', None, None)
            else:
                self.export_ui.switch_func(None, 'å¯¼å‡ºæŠ¥å‘Š', 'å…³é—­', None, None)
            # self.save_json(export_list)

    @staticmethod
    def save_json(export_list):
        f = asksaveasfile(mode='wb', defaultextension=".json")
        if f:
            json_str = json.dumps(export_list, ensure_ascii=False)
            f.write(json_str.encode(encoding='utf-8'))
            f.close()
            logger.debug('å¯¼å‡ºç»“æœä¿å­˜æˆåŠŸ')
            showinfo('å¯¼å‡º', 'å¯¼å‡ºæˆåŠŸï¼')
        else:
            logger.debug('ç”¨æˆ·å–æ¶ˆä¿å­˜å¯¼å‡ºç»“æœ')

    @staticmethod
    def get_new_path(i, seek, file_path, file_name_format, out_dir, cover='rename'):
        """
        è·å–ä¸€ä¸ªæ–°çš„æ–‡ä»¶è·¯å¾„
        :param cover: è¦†ç›–ç­–ç•¥ renameé‡å‘½åï¼ˆé»˜è®¤ï¼‰ï¼Œskipè·³è¿‡ï¼Œcoverè¦†ç›–
        :param i: æ–‡ä»¶ç¼–å·
        :param seek: å¯¼å‡ºé™„ä»¶ç¼–å·
        :param file_path: æ–‡ä»¶åï¼ˆè·¯å¾„ï¼‰
        :param out_dir: è¾“å‡ºç›®å½•
        :param file_name_format: æ–‡ä»¶åæ ¼å¼
        :return: æ–°çš„æ–‡ä»¶è·¯å¾„
        """
        if not os.path.isdir(out_dir):
            out_dir = os.path.dirname(os.path.abspath(out_dir))  # è·å–å¯¼å‡ºè·¯å¾„
        if file_name_format is None:
            file_name_format = '|è‡ªå¢ç¼–å·||è¿æ¥ç¬¦||åŸæ–‡ä»¶å||åç¼€å|'
        number = str(i).rjust(4, '0') + '.' + str(seek).rjust(2, '0')  # |è‡ªå¢ç¼–å·|ï¼Œä¸å¤Ÿå‰é¢æ·»0
        raw_file_name = os.path.basename(file_path)  # |åŸæ–‡ä»¶å|
        suffix = ''  # |åç¼€å|
        if '.' in raw_file_name:
            name_part = raw_file_name.split('.')
            part_count = len(name_part)
            suffix = name_part[-1]  # åç¼€å
            raw_file_name = '.'.join(name_part[:(part_count - 1)])  # å‰é¢çš„ï¼Œè¿™ä¹ˆå¤„ç†æ˜¯å› ä¸ºæœ‰çš„æ–‡ä»¶åä¸­æœ‰å¤šä¸ª .
        link_char = ' - '  # |è¿æ¥ç¬¦|
        file_name = out_dir + '/' + file_name_format.replace('|è‡ªå¢ç¼–å·|', number) \
            .replace('|è¿æ¥ç¬¦|', link_char) \
            .replace('|åŸæ–‡ä»¶å|', raw_file_name) \
            .replace('|åç¼€å|', '.' + suffix) \
            .strip()
        if 'cover' == cover:
            return file_name  # æ‹¼æ¥è·¯å¾„
        if 'skip' == cover:
            if os.path.isfile(file_name):
                return None
            else:
                return file_name
        if 'rename' == cover:
            add_index = 1
            while True:
                if not os.path.isfile(file_name):
                    return file_name
                file_name = out_dir + '/' + file_name_format.replace('|è‡ªå¢ç¼–å·|', number) \
                    .replace('|è¿æ¥ç¬¦|', link_char) \
                    .replace('|åŸæ–‡ä»¶å|', raw_file_name) \
                    .replace('|åç¼€å|', '.' + str(add_index) + '.' + suffix) \
                    .strip()
                add_index += 1

    @staticmethod
    def re_decode(s: str, encoding: str = 'gbk'):
        """
        é‡æ–°è§£ç ï¼Œè§£å†³oleobjå¯¹ä¸­æ–‡ä¹±ç çš„é—®é¢˜
        :param s: åŸå§‹å­—ç¬¦ä¸²
        :param encoding: æ–°çš„è§£ç ç¼–ç ï¼Œé»˜è®¤ä¸º GBK
        :return: æ–°çš„å­—ç¬¦ä¸²
        """
        i81 = s.encode('iso-8859-1')
        return i81.decode(encoding)

    def dispose(self, index: int, docx_file: str, parameter):
        """
        å¤„ç†docxæ–‡æ¡£
        :param index: ç´¢å¼•ï¼Œæ–‡æ¡£çº§
        :param docx_file: æ–‡æ¡£è·¯å¾„
        :param parameter: å…¶ä»–å‚æ•°
        :return: æ–‡æ¡£å¯¼å‡ºåˆ—è¡¨
        """
        save_way = parameter.get('way')
        output_dir = parameter.get('dir')
        name_format = parameter.get('name')
        export_type = parameter.get('type')
        is_del_raw = parameter.get('del')
        cover = parameter.get('cover')

        # å¯¼å‡ºè¿‡ç¨‹ä¸­æ˜¯å¦å‡ºé”™
        is_error = False

        # å¯¼å‡ºçš„æ–‡ä»¶åˆ—è¡¨
        export_files = []
        # ä»1å¼€å§‹
        index += 1

        # é™„ä»¶ç¼–å·
        seek = 0

        if not docx_file.endswith('.docx'):
            logger.debug('ä¸æ”¯æŒçš„æ–‡ä»¶ç±»å‹')
            return

            # å¯¼å‡ºæ–¹å¼è®¾ç½®
        if save_way == 1:
            output_dir = docx_file[:-5]
        elif save_way == 2:
            output_dir = os.path.dirname(docx_file)
        else:
            if output_dir:
                if not os.path.isdir(output_dir):
                    os.makedirs(output_dir)
            else:
                logger.error('å¯¼å‡ºè·¯å¾„ä¸æ­£ç¡®ï¼š' + docx_file)
                return export_files

        # å¯¼å‡ºæ–‡ä»¶å‰åå–å‡ºç©ºæ ¼ï¼ˆä¸å»é™¤å¯èƒ½å¯¼å‡ºå¤±è´¥ï¼‰
        output_dir = output_dir.strip()

        # åˆ›å»ºå¯¼å‡ºæ–‡ä»¶å¤¹
        if not os.path.isdir(output_dir):
            logger.debug('åˆ›å»ºå¯¼å‡ºæ–‡ä»¶å¤¹ï¼š%s', output_dir)
            os.mkdir(output_dir)
            if os.path.isdir(output_dir):
                logger.debug('åˆ›å»ºå¯¼å‡ºæ–‡ä»¶å¤¹æˆåŠŸ')
            else:
                logger.debug('åˆ›å»ºå¯¼å‡ºæ–‡ä»¶å¤¹å¤±è´¥')

        # æ‰“å¼€docxæ–‡æ¡£
        try:
            docx_document = docx.Document(docx_file)
        except Exception as e:
            logger.exception('æ‰“å¼€æ–‡æ¡£æ—¶å‡ºé”™%s' % str(e))
            self.remove_empty_dir(output_dir)
            return False
        logger.debug('æ‰“å¼€æ–‡æ¡£å®Œæˆ')

        # æ–‡æ¡£ä¿¡æ¯
        if 'info' in export_type:
            docx_properties = docx_document.core_properties
            all_properties = 'ä½œè€…\t' + str(docx_properties.author) + '\n'
            all_properties += 'ç±»åˆ«\t' + str(docx_properties.category) + '\n'
            all_properties += 'æ³¨é‡Š\t' + str(docx_properties.comments) + '\n'
            all_properties += 'å†…å®¹çŠ¶æ€\t' + str(docx_properties.content_status) + '\n'
            all_properties += 'åˆ›å»ºæ—¶é—´\t' + str(docx_properties.created) + '\n'
            all_properties += 'æ ‡è¯†ç¬¦\t' + str(docx_properties.identifier) + '\n'
            all_properties += 'å…³é”®å­—\t' + str(docx_properties.keywords) + '\n'
            all_properties += 'è¯­è¨€\t' + str(docx_properties.language) + '\n'
            all_properties += 'æœ€åä¿®æ”¹è€…\t' + str(docx_properties.last_modified_by) + '\n'
            all_properties += 'ä¸Šæ¬¡æ‰“å°\t' + str(docx_properties.last_printed) + '\n'
            all_properties += 'ä¿®æ”¹æ—¶é—´\t' + str(docx_properties.modified) + '\n'
            all_properties += 'ä¿®è®¢\t' + str(docx_properties.revision) + '\n'
            all_properties += 'ä¸»é¢˜\t' + str(docx_properties.subject) + '\n'
            all_properties += 'æ ‡é¢˜\t' + str(docx_properties.title) + '\n'
            all_properties += 'ç‰ˆæœ¬\t' + str(docx_properties.version) + '\n'
            logger.debug('æ–‡æ¡£ä¿¡æ¯ï¼š%s', all_properties.replace('\n', 'ï¼Œ '))

            # å¯¼å‡ºæ–‡æ¡£ä¿¡æ¯
            seek += 1
            info_file_path = self.get_new_path(index, seek, 'æ–‡æ¡£ä¿¡æ¯.txt', name_format, output_dir, cover)
            if info_file_path:
                logger.debug('æ–‡æ¡£ä¿¡æ¯ä¿å­˜ä½ç½®ï¼š%s', info_file_path)
                with open(info_file_path, 'w', encoding='utf-8') as f:
                    f.write(all_properties)
                export_files.append(info_file_path)
            else:
                logger.debug('è·³è¿‡æ–‡ä»¶')

        # æ‰€æœ‰æ–‡æœ¬
        all_text = ''
        if 'text' in export_type:
            for paragraph in docx_document.paragraphs:
                all_text += paragraph.text + ' '  # æ®µè½ä¹‹é—´ç”¨ç©ºæ ¼éš”å¼€
            logger.debug('æ‰€æœ‰æ–‡æœ¬ï¼š%s', all_text)

        # æ‰€æœ‰è¡¨æ ¼
        all_table_text = ''
        if 'table' in export_type:
            for table in docx_document.tables:
                for cell in getattr(table, '_cells'):
                    all_table_text += cell.text + '|'  # å•å…ƒæ ¼ä¹‹é—´ç”¨ â€œ|â€ éš”å¼€
            logger.debug('æ‰€æœ‰è¡¨æ ¼æ–‡æœ¬ï¼š%s', all_table_text)

        # å¯¼å‡ºæ–‡æœ¬
        if 'combine' in export_type and ('text' in export_type or 'table' in export_type):

            seek += 1
            combine_file_path = self.get_new_path(index, seek, 'æ–‡æœ¬å’Œè¡¨æ ¼.txt', name_format, output_dir,
                                                  cover)
            if combine_file_path:
                logger.debug('æ–‡æ¡£æ–‡æœ¬å’Œè¡¨æ ¼ä¿å­˜ä½ç½®ï¼š%s', combine_file_path)
                with open(combine_file_path, 'w', encoding='utf-8') as f:
                    f.write(all_text)
                    f.write('\n')  # æ¢ä¸ªè¡Œ
                    f.write(all_table_text)
                export_files.append(combine_file_path)
            else:
                logger.debug('è·³è¿‡æ–‡ä»¶')
        else:
            if 'text' in export_type:
                seek += 1
                text_file_path = self.get_new_path(index, seek, 'æ–‡æœ¬.txt', name_format, output_dir, cover)
                if text_file_path:
                    logger.debug('æ–‡æ¡£æ–‡æœ¬ä¿å­˜ä½ç½®ï¼š%s', text_file_path)
                    with open(text_file_path, 'w', encoding='utf-8') as f:
                        f.write(all_text)
                    export_files.append(text_file_path)
                else:
                    logger.debug('è·³è¿‡æ–‡ä»¶')

            if 'table' in export_type:
                seek += 1
                table_file_path = self.get_new_path(index, seek, 'è¡¨æ ¼.txt', name_format, output_dir, cover)
                if table_file_path:
                    logger.debug('æ–‡æ¡£è¡¨æ ¼ä¿å­˜ä½ç½®ï¼š%s', table_file_path)
                    with open(table_file_path, 'w', encoding='utf-8') as f:
                        f.write(all_table_text)
                    export_files.append(table_file_path)
                else:
                    logger.debug('è·³è¿‡æ–‡ä»¶')

        # éå†æ‰€æœ‰é™„ä»¶
        if 'image' in export_type or 'attachment' in export_type:
            docx_related_parts = docx_document.part.related_parts
            for part in docx_related_parts:
                part = docx_related_parts[part]
                part_name = str(part.partname)  # é™„ä»¶è·¯å¾„ï¼ˆpartnameï¼‰

                # åªå¯¼å‡ºè¿™ä¸¤ä¸ªç›®å½•ä¸‹çš„
                if not (part_name.startswith('/word/media/') or part_name.startswith('/word/embeddings/')):
                    continue

                # æ„å»ºå¯¼å‡ºè·¯å¾„
                seek += 1
                save_path = self.get_new_path(index, seek, part.partname, name_format, output_dir, cover)

                # ole æ–‡ä»¶åˆ¤æ–­
                # ä¸ç¬¦åˆ .bin ä½œä¸ºåç¼€ä¸”æ–‡ä»¶åä¸­æœ‰oleï¼Œåˆ™ä¸è¢«è®¤ä¸ºæ˜¯OLEæ–‡ä»¶
                if not (part_name.lower().endswith('.bin') and 'ole' in part_name.lower()):
                    # å¦‚æœæ²¡æœ‰æ”¯æŒå›¾ç‰‡å¯¼å‡º
                    if 'image' not in export_type:
                        continue

                    if save_path is None:
                        logger.debug('è·³è¿‡æ–‡ä»¶')
                        continue

                    # ç›´æ¥å†™å…¥æ–‡ä»¶
                    logger.debug('å›¾ç‰‡å¯¼å‡ºè·¯å¾„ï¼š%s', save_path)
                    with open(save_path, 'wb') as f:
                        f.write(part.blob)
                    export_files.append(save_path)  # è®°å½•æ–‡ä»¶

                    continue

                # å¦‚æœæ²¡æœ‰æ”¯æŒé™„ä»¶å¯¼å‡º
                if 'attachment' not in export_type:
                    continue

                # å°†å­—èŠ‚æ•°ç»„ä¼ é€’ç»™oleobjå¤„ç†
                for ole in oleobj.find_ole(save_path, part.blob):
                    if ole is None:  # æ²¡æœ‰æ‰¾åˆ° OLE æ–‡ä»¶ï¼Œè·³è¿‡
                        continue

                    for path_parts in ole.listdir():  # éå†OLEä¸­çš„æ–‡ä»¶

                        # åˆ¤æ–­æ˜¯ä¸æ˜¯[1]Ole10Nativeï¼Œä½¿ç”¨åˆ—è¡¨æ¨å¯¼å¼å¿½ç•¥å¤§å°å†™ï¼Œä¸æ˜¯çš„è¯å°±ä¸è¦ç»§ç»­äº†
                        if '\x01ole10native'.casefold() not in [path_part.casefold() for path_part in
                                                                path_parts]:
                            continue

                        stream = None
                        try:
                            # ä½¿ç”¨ Ole File æ‰“å¼€ OLE æ–‡ä»¶
                            stream = ole.openstream(path_parts)
                            opkg = oleobj.OleNativeStream(stream)
                        except IOError:
                            logger.debug('ä¸æ˜¯OLEæ–‡ä»¶ï¼š%s', path_parts)
                            if stream is not None:  # å…³é—­æ–‡ä»¶æµ
                                stream.close()
                            continue

                        # æ‰“å°ä¿¡æ¯
                        if opkg.is_link:
                            logger.debug('æ˜¯é“¾æ¥è€Œä¸æ˜¯æ–‡ä»¶ï¼Œè·³è¿‡')
                            continue

                        ole_filename = self.re_decode(opkg.filename)
                        ole_src_path = self.re_decode(opkg.src_path)
                        ole_temp_path = self.re_decode(opkg.temp_path)
                        logger.debug('æ–‡ä»¶åï¼š%sï¼ŒåŸè·¯å¾„ï¼š%sï¼Œç¼“å­˜è·¯å¾„ï¼š%s', ole_filename, ole_src_path, ole_temp_path)

                        # ç”Ÿæˆæ–°çš„æ–‡ä»¶å
                        seek += 1
                        filename = self.get_new_path(index, seek, ole_filename, name_format, output_dir, cover)

                        logger.debug('OLEé™„ä»¶å¯¼å‡ºè·¯å¾„ï¼š%s', filename)

                        if filename is None:
                            logger.debug('è·³è¿‡')
                            continue

                        # è½¬å­˜
                        try:
                            with open(filename, 'wb') as writer:
                                n_dumped = 0
                                next_size = min(oleobj.DUMP_CHUNK_SIZE, opkg.actual_size)
                                while next_size:
                                    data = stream.read(next_size)
                                    writer.write(data)
                                    n_dumped += len(data)
                                    if len(data) != next_size:
                                        logger.warning('æƒ³è¦è¯»å– %d, å®é™…å–å¾— %d', next_size, len(data))
                                        break
                                    next_size = min(oleobj.DUMP_CHUNK_SIZE, opkg.actual_size - n_dumped)
                            export_files.append(filename)  # è®°å½•å¯¼å‡ºçš„æ–‡ä»¶
                        except Exception as exc:
                            is_error = True
                            logger.exception('åœ¨è½¬å­˜æ—¶å‡ºç°é”™è¯¯', exc)
                        finally:
                            stream.close()
        logger.debug('å¯¼å‡ºçš„æ‰€æœ‰æ–‡ä»¶ï¼š%s', export_files)
        if not is_error and is_del_raw:
            logger.debug('åˆ é™¤åŸæ–‡ä»¶ï¼š%s', docx_file)
            os.remove(docx_file)
        self.remove_empty_dir(output_dir)
        return export_files

    @staticmethod
    def remove_empty_dir(target_dir):
        if os.path.isdir(target_dir):
            if not os.listdir(target_dir):
                logger.debug("åˆ é™¤ç©ºæ–‡ä»¶å¤¹ï¼š%s", target_dir)
                os.removedirs(target_dir)


def run():
    root = tk.Tk()
    Application(master=root, version=v)

    def close_window():
        ans = askyesno(title='æç¤º', message='æ˜¯å¦å…³é—­çª—å£ï¼Ÿ')
        if ans:
            root.destroy()
        else:
            return

    root.protocol('WM_DELETE_WINDOW', close_window)

    root.title('å¯¼å‡ºdocx')
    root.iconbitmap(resource_path('images/icon.ico'))
    root.mainloop()


if __name__ == '__main__':
    run()
